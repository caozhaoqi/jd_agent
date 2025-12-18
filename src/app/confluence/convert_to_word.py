import os
import json
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from app.utils.logger import logger

# 配置路径
CURRENT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.dirname(CURRENT_DIR)))
DATA_DIR = os.path.join(PROJECT_ROOT, "confluence_data")
OUTPUT_DIR = os.path.join(PROJECT_ROOT, "confluence_word_documents")
IMAGE_DIR = os.path.join(OUTPUT_DIR, "images")

# 创建输出目录
os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(IMAGE_DIR, exist_ok=True)


def download_image(url, page_id, image_index):
    """下载图片并保存到本地"""
    try:
        response = requests.get(url, timeout=10)
        if response.status_code == 200:
            # 检查Content-Type是否为图片
            content_type = response.headers.get("Content-Type", "")
            if not content_type.startswith("image/"):
                logger.warning(
                    f"⚠️  不是图片类型，跳过: {url} (Content-Type: {content_type})"
                )
                return None

            # 生成图片文件名
            image_ext = (
                content_type.split("/")[-1] if content_type else url.split(".")[-1]
            )
            if image_ext in ["png", "jpg", "jpeg", "gif", "bmp"]:
                pass
            else:
                image_ext = "png"  # 默认使用png格式
            image_filename = f"page_{page_id}_image_{image_index}.{image_ext}"
            image_path = os.path.join(IMAGE_DIR, image_filename)

            # 保存图片
            with open(image_path, "wb") as f:
                f.write(response.content)

            logger.info(f"✅ 下载图片成功: {image_filename}")
            return image_path
    except Exception as e:
        logger.error(f"❌ 下载图片失败 {url}: {str(e)}")
    return None


def convert_confluence_html_to_docx(html_content, doc):
    """将Confluence HTML内容转换为Word文档"""
    soup = BeautifulSoup(html_content, "html.parser")

    # 处理段落
    paragraphs = soup.find_all(["p", "h1", "h2", "h3", "h4", "h5", "h6"])
    for para in paragraphs:
        # 处理标题
        if para.name.startswith("h"):
            level = int(para.name[1])
            if level == 1:
                doc.add_heading(para.get_text(strip=True), level=0)
            else:
                doc.add_heading(para.get_text(strip=True), level=level - 1)
        else:
            # 处理普通段落
            doc.add_paragraph(para.get_text(strip=True))

    # 处理图片 - 支持标准img标签和Confluence特有ac:image标签
    # 先处理标准img标签
    images = soup.find_all("img")
    for i, img in enumerate(images):
        img_url = img.get("src")
        if img_url:
            # 下载图片
            image_path = download_image(img_url, page_id, i)
            if image_path:
                # 添加图片到文档（添加错误处理）
                try:
                    doc.add_picture(image_path, width=Inches(5))
                except Exception as e:
                    logger.error(f"❌ 添加图片到文档失败 {image_path}: {str(e)}")

    # 处理Confluence特有ac:image标签
    ac_images = soup.find_all("ac:image")
    for i, ac_img in enumerate(ac_images):
        # 获取附件文件名
        ri_attachment = ac_img.find("ri:attachment")
        if ri_attachment:
            filename = ri_attachment.get("ri:filename")
            if filename:
                # 构建图片URL (Confluence附件URL格式)
                # 注意：这里需要根据实际Confluence服务器地址调整
                img_url = f"https://wiki.hcmcloud.cn/download/attachments/{page_id}/{filename}"
                # 下载图片
                image_path = download_image(img_url, page_id, len(images) + i)
                if image_path:
                    # 添加图片到文档（添加错误处理）
                    try:
                        doc.add_picture(image_path, width=Inches(5))
                    except Exception as e:
                        logger.error(f"❌ 添加图片到文档失败 {image_path}: {str(e)}")

    # 处理列表
    lists = soup.find_all(["ul", "ol"])
    for lst in lists:
        list_items = lst.find_all("li")
        if list_items:
            # 创建一个新段落
            p = doc.add_paragraph()
            for item in list_items:
                p.add_run(item.get_text(strip=True)).bold = True
                p.add_run("\n")


def convert_confluence_pages_to_word():
    """将Confluence页面数据转换为Word文档"""
    logger.info("🚀 开始将Confluence页面转换为Word文档...")

    # 读取页面数据
    pages_file = os.path.join(DATA_DIR, "confluence_pages.json")
    if not os.path.exists(pages_file):
        logger.error("❌ 未找到Confluence页面数据文件")
        return

    with open(pages_file, "r", encoding="utf-8") as f:
        pages = json.load(f)

    logger.info(f"✅ 成功加载 {len(pages)} 个页面数据")

    # 转换每个页面
    for i, page in enumerate(pages, 1):
        global page_id  # 使用全局变量传递page_id给download_image函数
        page_id = page["page_id"]

        logger.info(f"🔄 正在转换页面 {i}/{len(pages)}: {page['title']}")

        # 创建新的Word文档
        doc = Document()

        # 添加页面标题
        doc.add_heading(page["title"], level=0)

        # 添加页面信息
        doc.add_paragraph(f"URL: {page['url']}")
        doc.add_paragraph(f"作者: {page['author']}")
        doc.add_paragraph(f"创建时间: {page['created_at']}")
        doc.add_paragraph(f"更新时间: {page['updated_at']}")

        # 添加分隔线
        doc.add_paragraph("=" * 50)

        # 转换内容
        if page["content"]:
            convert_confluence_html_to_docx(page["content"], doc)

        # 保存文档
        doc_filename = f"{page['title']}.docx".replace("/", "_").replace("\\", "_")
        doc_path = os.path.join(OUTPUT_DIR, doc_filename)

        try:
            doc.save(doc_path)
            logger.info(f"✅ 保存Word文档成功: {doc_filename}")
        except Exception as e:
            logger.error(f"❌ 保存Word文档失败 {doc_filename}: {str(e)}")

    logger.info(f"🎉 所有页面转换完成！文档保存在: {OUTPUT_DIR}")


if __name__ == "__main__":
    convert_confluence_pages_to_word()
